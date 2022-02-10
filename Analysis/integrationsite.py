import glob, os
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import seaborn as sns
from docx import Document
from docx.shared import Cm, Inches
import dataframe_image as dfi
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('Integration site analysis_format.docx')
writer = pd.ExcelWriter('Integration site analysis.xlsx', engine='xlsxwriter')
def summary():
	dic={}
	dic2={}
	for fl in glob.glob("*.final.xlsx"):
		fd=pd.read_excel(fl)
		sn=fl.split(".")[0]
		dic[sn]=[len(fd),len(fd[fd["Func.refGene"]=="exonic"]),len(fd[fd["Func.refGene"]=="intronic"]),len(fd[fd["Func.refGene"]=="intergenic"]),len(fd[fd["Func.refGene"]=="splicing"]),len(fd[fd["Func.refGene"]=="exonic;splicing"]),len(fd[fd["Func.refGene"]=="upstream"]),len(fd[fd["Func.refGene"]=="downstream"]),len(fd[fd["Func.refGene"]=="upstream;downstream"]),len(fd[fd["Func.refGene"]=="ncRNA_exonic"]),len(fd[fd["Func.refGene"]=="ncRNA_exonic;splicing"]),len(fd[fd["Func.refGene"]=="ncRNA_intronic"]),len(fd[fd["Func.refGene"]=="ncRNA_splicing"]),len(fd[fd["Func.refGene"]=="ncRNA_UTR5"]),len(fd[fd["Func.refGene"]=="ncRNA_UTR3"]),len(fd[fd["Func.refGene"]=="UTR3"]),len(fd[fd["Func.refGene"]=="UTR5"]),len(fd[fd["Func.refGene"]=="UTR5;UTR3"])]
	df=pd.DataFrame(dic)
	df.insert(0,"Func.refGene",['Total','Exonic','Intronic','Intergenic','Splicing','Exonic;splicing','Upstream','Downstream','Upstream;downstream','ncRNA_exonic','ncRNA_exonic;splicing','ncRNA_intronic','ncRNA_splicing','ncRNA_UTR5','ncRNA_UTR3','UTR3','UTR5','UTR5;UTR3'])
	dfstyle=df.style.set_properties(subset=['Func.refGene'],**{'font-weight':'bold'})
	dfstyle.to_excel(writer,sheet_name='integration sites',index=0)
	worksheet = writer.sheets['integration sites']
	worksheet.set_column('A:A',22)
	df2=df.drop(df.index[0])
	legends=list(df2['Func.refGene'])
	df2=df2.drop(['Func.refGene'],axis='columns')
	df2.plot.pie(subplots=True,figsize=(13,4),legend=None,labels=None,fontsize=7,colors=sns.color_palette('tab20',len(df2.index)),counterclock=False, startangle=90)
	plt.legend(legends,loc='center right',ncol=9,fontsize=7, bbox_to_anchor=(1,0))
	plt.suptitle('Distribution of integration sites')
	plt.savefig("summary_plot.png", bbox_inches='tight')
	plt.close()
	df=df.set_index("Func.refGene")
	dfi.export(df,'summary_table.png')
	la2=doc.paragraphs[8]
	la2.add_run().add_picture("summary_table.png")
	la=doc.paragraphs[10]
	la.add_run().add_picture("summary_plot.png", width=Cm(15), height=Cm(6))
	la.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
def TSSd():
	tss_df=pd.DataFrame(columns=['Sample ID','d < 2.5kb', 'd < 5kb', 'd < 10kb'])
	for tss_fl in glob.glob("*.final.xlsx"):
		tss_sn=tss_fl.split(".")[0]
		tss_fd=pd.read_excel(tss_fl)
		tss_fd=tss_fd[~tss_fd["Distance2TSS"].astype(str).str.contains("\.",na=False)]
		tss_fd["Distance2TSS"]=pd.to_numeric(tss_fd["Distance2TSS"])
		tss_df=tss_df.append({'Sample ID':tss_sn, 'd < 2.5kb':len(tss_fd[tss_fd["Distance2TSS"]<2500]), 'd < 5kb' : len(tss_fd[tss_fd["Distance2TSS"]<5000]) , 'd < 10kb' : len(tss_fd[tss_fd["Distance2TSS"]<10000])},ignore_index=True)
		u=sns.kdeplot(data=tss_fd,x=tss_fd['Distance2TSS'],linewidth=0.5)
		plt.xlim(0,500000)
		xlabels = ['{:,.1f}'.format(x) + 'kb' for x in u.get_xticks()/1000]
		ylabels = ['{:,.1f}'.format(x) for x in u.get_yticks()*100000]
		u.set_xticklabels(xlabels,fontsize = 7)
		u.set_yticklabels(ylabels,fontsize = 7)
	tss_dfstyle=tss_df.style.set_properties(subset=['Sample ID'],**{'font-weight':'bold'})
	tss_dfstyle.to_excel(writer,sheet_name='TSS distance',index=0)
	worksheet = writer.sheets['TSS distance']
	worksheet.set_column('A:D',10)
	plt.legend(tss_df['Sample ID'])
	plt.title('Distance distribution between integration sites and TSSs',fontsize=9)
	plt.xlabel('Distance',fontsize=8)
	plt.ylabel('Density(*1e-5)',fontsize=8)
	sns.set(rc={'figure.figsize':(5,4)})
	plt.savefig("TSS_plot.png", bbox_inches='tight')
	sns.reset_orig()
	plt.close()
	tss_df.index=tss_df.index+1
	dfi.export(tss_df,'TSS_table.png')
	tt=doc.paragraphs[16]
	tt.add_run().add_picture("TSS_table.png")
	tla=doc.paragraphs[18]
	tla.add_run().add_picture("TSS_plot.png", width=Cm(10), height=Cm(7))
	tla.alignment = WD_ALIGN_PARAGRAPH.CENTER

def CpGd():
	cpg_df=pd.DataFrame(columns=['Sample ID','Intragenic','d < 2.5kb', 'd < 5kb', 'd < 10kb'])
	for cpg_fl in glob.glob("*.final.xlsx"):
		cpg_sn=cpg_fl.split(".")[0]
		cpg_fd=pd.read_excel(cpg_fl)
		cpg_fd=cpg_fd[~cpg_fd["Distance2CpG"].astype(str).str.contains("\.",na=False)]
		cpg_fd=cpg_fd[(cpg_fd["Func.refGene"]=="exonic") | (cpg_fd["Func.refGene"]=="intronic")]
		cpg_fd["Distance2CpG"]=pd.to_numeric(cpg_fd["Distance2CpG"])
		cpg_df=cpg_df.append({'Sample ID':cpg_sn, 'Intragenic':len(cpg_fd),'d < 2.5kb':len(cpg_fd[cpg_fd["Distance2CpG"]<2500]), 'd < 5kb' : len(cpg_fd[cpg_fd["Distance2CpG"]<5000]) , 'd < 10kb' : len(cpg_fd[cpg_fd["Distance2CpG"]<10000])},ignore_index=True)
		cu=sns.kdeplot(data=cpg_fd,x=cpg_fd['Distance2CpG'],linewidth=0.5)
		plt.xlim(0,500000)
		cxlabels = ['{:,.1f}'.format(x) + 'kb' for x in cu.get_xticks()/1000]
		cylabels = ['{:,.1f}'.format(x) for x in cu.get_yticks()*100000]
		cu.set_xticklabels(cxlabels,fontsize = 7)
		cu.set_yticklabels(cylabels,fontsize = 7)
	cpg_dfstyle=cpg_df.style.set_properties(subset=['Sample ID'],**{'font-weight':'bold'})
	cpg_dfstyle.to_excel(writer,sheet_name='CpG distance',index=0)
	worksheet = writer.sheets['CpG distance']
	worksheet.set_column('A:E',10)
	plt.legend(cpg_df['Sample ID'])
	plt.title('Distance distribution between integration sites and CpG islands',fontsize=9)
	plt.xlabel('Distance',fontsize=8)
	plt.ylabel('Density(*1e-5)',fontsize=8)
	sns.set(rc={'figure.figsize':(5,4)})
	plt.savefig("CpG_plot.png", bbox_inches='tight')
	sns.reset_orig()
	plt.close()
	cpg_df.index=cpg_df.index+1
	dfi.export(cpg_df,'CpG_table.png')
	cla2=doc.paragraphs[27]
	cla2.add_run().add_picture("CpG_table.png")
	cla=doc.paragraphs[29]
	cla.add_run().add_picture("CpG_plot.png", width=Cm(10), height=Cm(7))
	cla.alignment = WD_ALIGN_PARAGRAPH.CENTER

def OCGTSG():
	df=pd.DataFrame(columns=['Sample ID','Number of TSGs','TSGs list', 'Number of OCGs', 'OCGs list'])
	for fl in glob.glob("*.final.xlsx"):
		ocg=pd.read_csv("OCG_CancerQuest_OCGTSG_220103_trim.tsv",sep='\t')
		ocgl=ocg["OCG"].tolist()
		tsg=pd.read_csv("TSG_CancerQuest_OCGTSG_220103_trim.tsv",sep='\t')
		tsgl=tsg["TSG"].tolist()
		sn=fl.split(".")[0]
		fd=pd.read_excel(fl)
		fd=fd[fd["Func.refGene"]=="exonic"]	
		lll=pd.DataFrame(columns=fd.columns)
		lll2=pd.DataFrame(columns=fd.columns)
		for i in ocgl:
			lll=lll.append(fd[fd["Gene1 name"]==i],ignore_index=True)
			lll=lll.append(fd[fd["Gene2 name"]==i],ignore_index=True)
		lll=lll.drop_duplicates()
		llls=lll["Gene.refGene"].tolist()
		setl=set(llls)	
		nll=list(setl)
		for i in tsgl:
			lll2=lll2.append(fd[fd["Gene1 name"]==i],ignore_index=True)
			lll2=lll2.append(fd[fd["Gene2 name"]==i],ignore_index=True)
		lll2=lll2.drop_duplicates()
		llls2=lll2["Gene.refGene"].tolist()
		setl2=set(llls2)	
		nll2=list(setl2)
		df=df.append({'Sample ID':sn, 'Number of TSGs':len(nll2),'TSGs list':nll2,'Number of OCGs':len(nll),'OCGs list':nll},ignore_index=True)
	df.to_excel(writer,sheet_name='OCG,TSG list',index=0)
	ola=doc.paragraphs[22]
	df=df.drop(['TSGs list','OCGs list'],axis='columns')
	df.index=df.index+1
	dfi.export(df,'OCGTSG_table.png')
	ola.add_run().add_picture("OCGTSG_table.png")


summary()
TSSd()
CpGd()
OCGTSG()

writer.save()
doc.save("Integration site analysis(primary)_220110.docx")
